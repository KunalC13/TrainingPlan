from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.llms import Ollama
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
import pandas as pd
from docx import Document  # New import to read DOCX files

# Step 1: Load Pre-trained Sentence Transformer Model
embedding_model = HuggingFaceEmbeddings(model_name='sentence-transformers/all-MiniLM-L6-v2')

# Step 2: Load Rules from a Document File
# Function to load rules from a Markdown file
def load_rules_from_markdown(filepath):
    """
    Reads a MD file and returns a list of non-empty paragraphs.
    Each paragraph is considered as one rule.
    """
    with open(filepath, 'r') as file:
        lines = file.readlines()
    rules = [line.strip() for line in lines if line.strip() and not line.startswith('#')]
    return rules

# Function to load rules from a DOCX file using python-docx
def load_rules_from_docx(filepath):
    """
    Reads a DOCX file and returns a list of non-empty paragraphs.
    Each paragraph is considered as one rule.
    """
    document = Document(filepath)
    rules = [para.text.strip() for para in document.paragraphs if para.text.strip()]
    return rules

use_docx = True  # Set to False if you prefer using a Markdown file

if use_docx:
    # Load the rule book from the DOCX file located in the "data" directory
    rules = load_rules_from_docx('data/Rule Book for Training Plans.docx')
else:
    # Load rules from the Markdown file (ensure the file exists in your data directory)
    rules = load_rules_from_markdown('data/rules.md')

# Load the rule book from the DOCX file located in the "data" directory
rules = load_rules_from_docx('data/Rule Book for Training Plans.docx')

# Optionally, if you want to load the exercise library for additional context, you can do so:
def load_text_from_docx(filepath):
    document = Document(filepath)
    return "\n".join([para.text.strip() for para in document.paragraphs if para.text.strip()])

# Uncomment the following line if you wish to use the exercise library as well
# exercise_library_text = load_text_from_docx('data/Exercise library App.docx')

# Step 3: Create FAISS Vector Store using from_texts
faiss_index = FAISS.from_texts(rules, embedding_model)

# Step 4: Set up Ollama LLaMA 3.2 with LangChain
llama_llm = Ollama(model='llama3.2:latest')

# Step 5: Create Prompt Template
prompt_template = PromptTemplate(
    input_variables=['retrieved_rules', 'user_profile'],
    template="""
    You are a professional fitness coach. Based on the following workout rules:
    {retrieved_rules}

    And the user profile:
    {user_profile}

    Generate a personalized, structured workout plan that adheres to the rules and fits the user's profile.
    """
)

# Step 6: Create LLM Chain
chain = LLMChain(llm=llama_llm, prompt=prompt_template)

# Step 7: Format User Profile
def format_user_profile(user_name, user_data_row):
    profile = f"""
    **User Profile:**
    - Name: {user_name}
    - Age: {user_data_row.get("Age", "N/A")}
    - Gender: {user_data_row.get("Gender", "N/A")}
    - Height: {user_data_row.get("Height (cm)", "N/A")} cm
    - Weight: {user_data_row.get("Weight (Kg)", "N/A")} kg
    - Fitness Goal(s): {user_data_row.get("What are your fitness goals? (You can select multiple options)", "N/A")}
    - Fitness Level: {user_data_row.get("How would you rate your fitness level?", "N/A")}
    - Workout Frequency: {user_data_row.get("How many days per week are you planning to working out?", "N/A")} days per week
    - Workout Duration: {user_data_row.get("How much time can you dedicate to each workout session?", "N/A")}
    - Workout Location: {user_data_row.get("Where are you planning to work out?", "N/A")}
    - Equipment (if home): {user_data_row.get("if you answered Home or mix, what equipment do you have available?", "N/A")}
    - Focus Areas: {user_data_row.get("Do you have any specific areas of the body you want to focus on?", "N/A")}
    - Cardio Preference: {user_data_row.get("Would you like to include cardio in your fitness program?", "N/A")}
    - Daily Activity Level: {user_data_row.get("What is your daily activity?", "N/A")}
    - Injury/Medical Condition: {user_data_row.get("Do you have any injury or medical condition?", "N/A")}
    - Injury Details (if any): {user_data_row.get("If you answered yes to above question, please provide a short description.", "N/A")}
    """
    return profile

# Step 8: Generate Workout Plan Based on User Profile
def generate_workout_plan(user_name, user_data_row):
    user_profile = format_user_profile(user_name, user_data_row)

    # Use profile information to guide retrieval
    docs = faiss_index.similarity_search(user_profile, k=5)
    retrieved_rules = "\n".join([doc.page_content for doc in docs])

    # Run the chain to generate workout
    response = chain.run({"retrieved_rules": retrieved_rules, "user_profile": user_profile})
    return response

# Step 9: Main Function
def main():
    input_csv = "responses_n.csv"  # Change this to your actual file path
    output_csv = "workout_plans1.csv"
    
    df = pd.read_csv(input_csv)
    
    # Generate workout plans for each user
    workout_plans = []
    for _, row in df.iterrows():
        user_name = row["Name "]
        user_data_row = {
            "Age": row["Age"],
            "Gender": row["Gender"],
            "Height (cm)": row["Height (cm)"],
            "Weight (Kg)": row["Weight (Kg)"],
            "What are your fitness goals? (You can select multiple options) ": row["What are your fitness goals? (You can select multiple options) "],
            "How would you rate your fitness level?": row["How would you rate your fitness level?"],
            "How many days per week are you planning to working out?": row["How many days per week are you planning to working out?"],
            "How much time can you dedicate to each workout session?": row["How much time can you dedicate to each workout session?"],
            "Where are you planning to work out?": row["Where are you planning to work out?"],
            "if you answered Home or mix, what equipment do you have available?": row["if you answered Home or mix, what equipment do you have available?"],
            "Do you have any specific areas of the body you want to focus on?": row["Do you have any specific areas of the body you want to focus on?"],
            "Would you like to include cardio in your fitness program?": row["Would you like to include cardio in your fitness program?If yes, please select the type(s) of cardio you are interested in:"],
            "What is your daily activity?": row["What is your daily activity?"],
            "Do you have any injury or medical condition?": row["Do you have any injury or medical condition?"],
            "If you answered yes to above question, please provide a short description.": row["If you answered yes to above question, please provide a short description."],
        }
        workout_plan = generate_workout_plan(user_name, user_data_row)
        workout_plans.append(workout_plan)
    
    df["Training_Plan"] = workout_plans
    
    # Save the updated dataframe with workout plans to a new CSV
    df.to_csv(output_csv, index=False)
    print(f"Workout plans saved to {output_csv}")

if __name__ == "__main__":
    main()
