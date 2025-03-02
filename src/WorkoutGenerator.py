# workout_generator.py
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.llms import Ollama
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain.output_parsers import PydanticOutputParser
from docx import Document
from pydantic import BaseModel
from typing import List, Optional
import json

# Define Pydantic models for the structured output
class Exercise(BaseModel):
    Exercise: str
    Sets: str
    Reps: str
    Notes: Optional[str] = ""

class DayPlan(BaseModel):
    Day: str
    Focus: str
    WarmUp: List[str]
    MainExercises: List[Exercise]
    CoolDown: List[str]
    AdditionalNotes: Optional[str] = ""

class WorkoutPlan(BaseModel):
    WorkoutPlan: List[DayPlan]

class WorkoutGenerator:
    def __init__(
        self,
        use_docx: bool = True,
        rule_file_path: str = 'data/RuleBook.docx',
        markdown_rule_path: str = 'data/rules.md',
        exercise_lib_path: str = 'data/ExerciseLib.docx'
    ):
        # Initialize the embedding model
        self.embedding_model = HuggingFaceEmbeddings(model_name='sentence-transformers/all-MiniLM-L6-v2')
        self.use_docx = use_docx

        # Load rules from DOCX or Markdown
        if self.use_docx:
            self.rules = self.load_rules_from_docx(rule_file_path)
        else:
            self.rules = self.load_rules_from_markdown(markdown_rule_path)
        
        # Load the exercise library text
        self.exercise_library_text = self.load_text_from_docx(exercise_lib_path)
        
        # Create FAISS index from the rules
        self.faiss_index = FAISS.from_texts(self.rules, self.embedding_model)
        
        # Set up the LLM (ensure your local Ollama service is running)
        self.llama_llm = Ollama(model='llama3.2:latest')
        
        # Create a PydanticOutputParser from our WorkoutPlan model.
        self.output_parser = PydanticOutputParser(pydantic_object=WorkoutPlan)
        format_instructions = self.output_parser.get_format_instructions()

        # Include "format_instructions" as an input variable so the parser's instructions aren’t inlined.
        self.prompt_template = PromptTemplate(
            input_variables=['retrieved_rules', 'exercise_library_text', 'user_profile', 'format_instructions'],
            template="""
                You are a professional fitness coach. Based on the following workout rules:
                {retrieved_rules}

                The following is the available exercise library:
                {exercise_library_text}

                And the user profile:
                {user_profile}

                Generate a personalized, structured workout plan that adheres to the rules and selects appropriate exercises from the exercise library. The exercises given should be aptly according to the workout duration of the user.
                Please format your answer according to these instructions exactly:
                {format_instructions}
                """
        )
        self.chain = LLMChain(llm=self.llama_llm, prompt=self.prompt_template)

    def load_rules_from_markdown(self, filepath: str) -> List[str]:
        """Reads a Markdown file and returns a list of non-empty paragraphs."""
        with open(filepath, 'r') as file:
            lines = file.readlines()
        return [line.strip() for line in lines if line.strip() and not line.startswith('#')]

    def load_rules_from_docx(self, filepath: str) -> List[str]:
        """Reads a DOCX file and returns a list of non-empty paragraphs."""
        document = Document(filepath)
        return [para.text.strip() for para in document.paragraphs if para.text.strip()]

    def load_text_from_docx(self, filepath: str) -> str:
        """Reads a DOCX file and returns its full text (non-empty paragraphs separated by newlines)."""
        document = Document(filepath)
        return "\n".join([para.text.strip() for para in document.paragraphs if para.text.strip()])

    def _format_user_profile(self, user_name: str, user_data: dict) -> str:
        """Formats the user profile information into a single string."""
        return f"""
            **User Profile:**
            - Name: {user_name}
            - Age: {user_data.get("Age", "N/A")}
            - Gender: {user_data.get("Gender", "N/A")}
            - Height: {user_data.get("Height (cm)", "N/A")} cm
            - Weight: {user_data.get("Weight (Kg)", "N/A")} kg
            - Fitness Goals: {user_data.get("Fitness Goals", "N/A")}
            - Fitness Level: {user_data.get("Fitness Level", "N/A")}
            - Workout Frequency: {user_data.get("Workout Frequency", "N/A")} days per week
            - Workout Duration: {user_data.get("Workout Duration", "N/A")}
            - Workout Location: {user_data.get("Workout Location", "N/A")}
            - Equipment: {user_data.get("Equipment", "N/A")}
            - Focus Areas: {user_data.get("Focus Areas", "N/A")}
            - Cardio Preference: {user_data.get("Cardio Preference", "N/A")}
            - Daily Activity: {user_data.get("Daily Activity", "N/A")}
            - Injury/Medical Condition: {user_data.get("Injury", "N/A")}
            - Injury Details: {user_data.get("Injury Details", "N/A")}
            """

    def generate_workout_plan(self, user_name: str, user_data: dict) -> WorkoutPlan:
        """Generates and returns a workout plan for the given user as a WorkoutPlan model."""
        user_profile = self._format_user_profile(user_name, user_data)
        # Retrieve the most relevant rules based on the user profile
        docs = self.faiss_index.similarity_search(user_profile, k=5)
        retrieved_rules = "\n".join([doc.page_content for doc in docs])
        
        # Run the chain to generate a response
        llm_output = self.chain.run({
            "retrieved_rules": retrieved_rules,
            "exercise_library_text": self.exercise_library_text,
            "user_profile": user_profile,
            "format_instructions": self.output_parser.get_format_instructions()
        })
        
        try:
            parsed_plan = self.output_parser.parse(llm_output)
        except Exception as e:
            print("Parsing failed:", e)
            parsed_plan = None
        return parsed_plan

    def display_workout_plan(self, workout_plan: WorkoutPlan):
        """Displays the workout plan in a structured format."""
        for day in workout_plan.WorkoutPlan:
            print(f"Day: {day.Day}")
            print(f"Focus: {day.Focus}")
            print("Warm-Up:")
            for exercise in day.WarmUp:
                print(f" - {exercise}")
            print("Main Exercises:")
            for ex in day.MainExercises:
                print(f" - {ex.Exercise} | Sets: {ex.Sets}, Reps: {ex.Reps} | Notes: {ex.Notes}")
            print("Cool Down:")
            for exercise in day.CoolDown:
                print(f" - {exercise}")
            if day.AdditionalNotes:
                print("Additional Notes:", day.AdditionalNotes)
            print("\n" + "-" * 30 + "\n")
